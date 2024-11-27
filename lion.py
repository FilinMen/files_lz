import docx 
from collections import Counter
import re #(регулярные выражения) модуль для работы с текстом 
import pandas as pd

doc = docx.Document('lion.docx')

num_par = len(doc.paragraphs) # количество параграфов
word_list = [] #список для слов без повторений 

all_letters = []
rus_let = re.compile(r'[а-яё]', re.IGNORECASE) #регулярное выражение для поиска русских букв

rus_text = [] #список для всего текста (только русские слова)
rus_word = re.compile(r'\b[а-яё]+\b',re.IGNORECASE) #регулярное выражение для поиска русских слов 

for paragraph in doc.paragraphs: #проходимся по всем параграфам в документе
    text = paragraph.text.lower() # приводим текст к нижнему регистру
    words = rus_word.findall(text) # находим все русские слова
    letters = rus_let.findall(text) # находим все русские буквы
    all_letters.extend(letters) #добовляем буквы в общий список

    for word in words:
        rus_text.append(word)
        if word not in word_list: # проверка на на наличие слова в списке 
            word_list.append(word) #добавления слова 

#---------------------------------------------------------------------------------------
#ВСТРЕЧАЕМОСТЬ РУССКИХ СЛОВ В ТЕКСТЕ

word_counts = Counter(rus_text) #подсчитываем частоту встречаемости слов
df_word = pd.DataFrame(word_counts.items(),columns=['Word','Quantity'])

doc_word = docx.Document() #создаю новый документ

doc_word.add_heading('Встречаемость слов в тексте', level=1)#добовляем заголовок

# Добавляем таблицу
table = doc_word.add_table(rows=1, cols=len(df_word.columns))

#Добавляем заголовки колонок
hdr_cells = table.rows[0].cells
for i, column_name in enumerate(df_word.columns):
    hdr_cells[i].text = column_name

# Добавляем данные из DataFrame
for index, row in df_word.iterrows():
    row_cells = table.add_row().cells
    for i, cell_value in enumerate(row):
         row_cells[i].text = str(cell_value)
         
doc_word.save('встречаемость_слов.docx') # Сохраняем документ

#--------------------------------------------------------------------------------------

#ВСТРЕЧАЕМОСТЬ БУКВ В ТЕКСТЕ
letters_counts = Counter(all_letters) #подсчитываем частоту встречаемости букв
df_letter = pd.DataFrame(letters_counts.items(),columns=['letter','Quantity'])

doc_letter = docx.Document() #создаю новый документ
doc_letter.add_heading('Встречаемость букв в тексте', level=1)#добовляем заголовок

# Добавляем таблицу
table = doc_letter.add_table(rows=1, cols=len(df_letter.columns))

#Добавляем заголовки колонок
hdr_cells = table.rows[0].cells
for i, column_name in enumerate(df_letter.columns):
    hdr_cells[i].text = column_name

# Добавляем данные из DataFrame
for index, row in df_letter.iterrows():
    row_cells = table.add_row().cells
    for i, cell_value in enumerate(row):
         row_cells[i].text = str(cell_value)

doc_letter.save('встречаемость_букв.docx') # Сохраняем документ

#--------------------------------------------------------------------------------------
print(len(word_list))
print(num_par)


